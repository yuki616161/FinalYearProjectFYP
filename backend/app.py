import os
import json
import time
import io
import docx
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from pypdf import PdfReader
from dotenv import load_dotenv

from google import genai
from google.genai import types

load_dotenv()

app = Flask(__name__)
CORS(app)

# --- GEMINI SETUP ---
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise RuntimeError("Missing GEMINI_API_KEY in .env")

client = genai.Client(api_key=api_key)

MODEL_CANDIDATES = ["gemini-2.5-flash-lite", "gemini-2.5-flash", "gemini-2.0-flash"]

# --- HELPERS ---
def extract_text_from_pdf(file_storage):
    try:
        reader = PdfReader(file_storage)
        text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"PDF Error: {e}")
        return None

def extract_text_from_docx(file_storage):
    try:
        doc = docx.Document(file_storage)
        return "\n".join([para.text for para in doc.paragraphs]).strip()
    except Exception as e:
        print(f"DOCX Error: {e}")
        return None

def call_gemini(prompt):
    response = None
    last_err = None
    for model_name in MODEL_CANDIDATES:
        for attempt in range(3):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                    config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0.3),
                )

                # Clean up potential markdown formatting before parsing
                raw_text = response.text.strip()
                if raw_text.startswith("```json"):
                    raw_text = raw_text[7:-3].strip()
                elif raw_text.startswith("```"):
                    raw_text = raw_text[3:-3].strip()

                return json.loads(raw_text), None
            except Exception as e:
                print(f"Gemini API Error ({model_name}, attempt {attempt+1}): {e}")
                last_err = e
                if "503" in str(e) or "overloaded" in str(e).lower():
                    time.sleep(1.5 * (attempt + 1))
                    continue
                if "404" in str(e):
                    break
    return None, last_err

def build_prompt(job_role, resume_text):
    return f"""
    You are an expert technical career coach and an advanced Applicant Tracking System (ATS).
    Analyze this resume for the target role: {job_role}.

    Task 1: Calculate the ATS Match Score (0-100) using a strict rubric: Hard Skills (40%), Experience (30%), Education (15%), Impact (15%).
    Task 2: Provide a 1-2 sentence explanation of WHY this score was given, referencing the rubric.
    Task 3: Identify key strengths and critical skill gaps.
    Task 4: Recommend a 3-step learning path to bridge the gaps. You MUST include a diverse mix of platforms (e.g., Coursera, Udemy, LinkedIn Learning, free YouTube tutorials, or official documentation).

    Resume Text:
    {resume_text}

    Return ONLY a JSON object with this exact structure (no markdown):
    {{
      "score": 85,
      "score_explanation": "The candidate scored highly in Hard Skills (35/40) and Education (15/15), but lost points in Impact (5/15) due to a lack of measurable achievements.",
      "feedback_text": "One concise paragraph of professional career advice.",
      "strengths": ["Strength 1", "Strength 2"],
      "gaps": ["Gap 1", "Gap 2"],
      "action_steps": ["Action 1", "Action 2"],
      "learning_path": [
        {{
          "skill": "React State Management",
          "course_name": "React Redux Full Crash Course",
          "platform": "YouTube",
          "difficulty": "Beginner",
          "estimated_time": "3 hours",
          "description": "A quick, free video tutorial to master global state management."
        }}
      ]
    }}
    """.strip()

# ==========================================
# ROUTE 1: FILE-BASED AI ANALYSIS
# ==========================================
@app.route("/api/analyze", methods=["POST"])
def analyze_resume():
    job_role = request.form.get("job_role", "").strip()
    if not job_role or "file" not in request.files:
        return jsonify({"error": "Missing job_role or file"}), 400

    file = request.files["file"]
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        resume_text = extract_text_from_pdf(file)
    elif filename.endswith(".docx"):
        resume_text = extract_text_from_docx(file)
    else:
        return jsonify({"error": "Invalid format."}), 400

    if not resume_text:
        return jsonify({"error": "Could not read text."}), 400

    prompt = build_prompt(job_role, resume_text[:30000])
    data, err = call_gemini(prompt)

    if err or not data:
        return jsonify({"error": "AI Service Busy. Please try again."}), 503
    return jsonify(data)

# ==========================================
# ROUTE 2: DIRECT AI ANALYSIS (NO FILE UPLOAD)
# ==========================================
@app.route("/api/analyze-direct", methods=["POST"])
def analyze_direct():
    data = request.json
    job_role = data.get("job_role", "").strip()
    resume_data = data.get("resume_data", {})

    if not job_role:
        return jsonify({"error": "Missing job_role"}), 400

    profile = resume_data.get('profile', {})
    resume_text = f"Name: {profile.get('name', '')}\nSummary: {profile.get('summary', '')}\n\n"

    resume_text += "Experience:\n"
    for exp in resume_data.get('experience', []):
        resume_text += f"- {exp.get('role')} at {exp.get('company')}: {exp.get('description')}\n"

    resume_text += "\nEducation:\n"
    for edu in resume_data.get('education', []):
        resume_text += f"- {edu.get('degree')} at {edu.get('school')}\n"

    resume_text += "\nProjects:\n"
    for proj in resume_data.get('projects', []):
        resume_text += f"- {proj.get('name')}: {proj.get('description')}\n"

    skills = [s.get('name', '') for s in resume_data.get('skills', [])]
    resume_text += f"\nSkills: {', '.join(skills)}\n"

    prompt = build_prompt(job_role, resume_text[:30000])
    result_data, err = call_gemini(prompt)

    if err or not result_data:
        return jsonify({"error": "AI Service Busy. Please try again."}), 503
    return jsonify(result_data)

# ==========================================
# ROUTE 3: RESUME BUILDER DOWNLOAD
# ==========================================
@app.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    data = request.json
    doc = docx.Document()

    profile = data.get('profile', {})
    doc.add_heading(profile.get('name', 'Your Name').upper(), level=0).alignment = 1

    contact = [profile.get(k) for k in ['email', 'phone', 'linkedin'] if profile.get(k)]
    if contact: doc.add_paragraph(" | ".join(contact)).alignment = 1
    if profile.get('summary'): doc.add_paragraph(profile.get('summary'))

    for section, title in [('experience', 'Experience'), ('education', 'Education'), ('projects', 'Projects')]:
        items = data.get(section, [])
        if items:
            doc.add_heading(title, level=1)
            for item in items:
                p = doc.add_paragraph()
                p.add_run(item.get('role', item.get('degree', item.get('name', '')))).bold = True
                p.add_run(f" - {item.get('company', item.get('school', item.get('link', '')))}")
                if item.get('description'): doc.add_paragraph(item.get('description')).style = 'List Bullet'

    for section, title in [('skills', 'Skills'), ('languages', 'Languages')]:
        items = data.get(section, [])
        if items:
            doc.add_heading(title, level=1)
            doc.add_paragraph(", ".join([i.get('name', '') for i in items if i.get('name')]))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    safe_name = profile.get('name', 'Resume').replace(' ', '_')
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"{safe_name}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    app.run(debug=True, port=5000)